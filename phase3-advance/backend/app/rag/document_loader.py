"""
Document Loader - Supports PDF, Word, Excel, CSV, TXT, Images (OCR).
Extracts text and prepares it for vector embedding.
"""
import os
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
from loguru import logger


class DocumentLoader:
    """Loads and extracts text from various document formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".txt",
        ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff",
    }

    def __init__(self):
        pass

    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document and extract its text content.

        Args:
            file_path: Path to the document file

        Returns:
            Dict with {filename, texts: [str, ...], metadata: {...}}
        """
        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}", "texts": []}

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return {"error": f"Unsupported format: {ext}", "texts": []}

        result = {
            "filename": path.name,
            "texts": [],
            "metadata": {"file_type": ext, "file_size": path.stat().st_size},
        }

        try:
            if ext == ".pdf":
                result["texts"] = self._load_pdf(file_path)
            elif ext in (".docx", ".doc"):
                result["texts"] = self._load_docx(file_path)
            elif ext in (".xlsx", ".xls"):
                result["texts"] = self._load_xlsx(file_path)
            elif ext == ".csv":
                result["texts"] = self._load_csv(file_path)
            elif ext == ".txt":
                result["texts"] = self._load_txt(file_path)
            elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff"):
                result["texts"] = self._load_image(file_path)

            logger.info(f"📄 Loaded {path.name}: {len(result['texts'])} text segments")
        except Exception as e:
            logger.error(f"Failed to load {path.name}: {e}")
            result["error"] = str(e)

        return result

    def _load_pdf(self, file_path: str) -> List[str]:
        """Extract text from PDF."""
        texts = []
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
            except ImportError:
                # Fallback: use pypdf
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
        except Exception as e:
            logger.warning(f"PDF loading degraded: {e}")

        if not texts:
            texts = [f"[PDF文件: {Path(file_path).name}]"]
        return texts

    def _load_docx(self, file_path: str) -> List[str]:
        """Extract text from Word document."""
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            if text and text.strip():
                return [text.strip()]
        except ImportError:
            pass

        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            if paragraphs:
                return ["\n".join(paragraphs)]
        except Exception as e:
            logger.warning(f"DOCX loading failed: {e}")

        return [f"[Word文档: {Path(file_path).name}]"]

    def _load_xlsx(self, file_path: str) -> List[str]:
        """Extract text from Excel spreadsheet."""
        try:
            import pandas as pd
            dfs = pd.read_excel(file_path, sheet_name=None)
            texts = []
            for sheet_name, df in dfs.items():
                text = f"工作表: {sheet_name}\n" + df.to_string(index=False)
                texts.append(text)
            return texts
        except Exception as e:
            logger.warning(f"XLSX loading failed: {e}")
            return [f"[Excel文件: {Path(file_path).name}]"]

    def _load_csv(self, file_path: str) -> List[str]:
        """Extract text from CSV."""
        try:
            import pandas as pd
            for encoding in ["utf-8", "gbk", "latin-1"]:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    return [df.to_string(index=False)]
                except:
                    continue
        except Exception as e:
            logger.warning(f"CSV loading failed: {e}")
        return [f"[CSV文件: {Path(file_path).name}]"]

    def _load_txt(self, file_path: str) -> List[str]:
        """Extract text from plain text file."""
        for encoding in ["utf-8", "gbk", "latin-1"]:
            try:
                text = Path(file_path).read_text(encoding=encoding)
                if text.strip():
                    return [text.strip()]
            except:
                continue
        return [f"[文本文件: {Path(file_path).name}]"]

    def _load_image(self, file_path: str) -> List[str]:
        """Extract text from image using OCR."""
        try:
            from PIL import Image
            import pytesseract
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            if text and text.strip():
                return [text.strip()]
        except ImportError:
            logger.warning("pytesseract not available for OCR")
        except Exception as e:
            logger.warning(f"Image OCR failed: {e}")
        return [f"[图片文件: {Path(file_path).name}]"]


def load_and_prepare(file_path: str) -> Dict[str, Any]:
    """Convenience function: load document and return texts ready for embedding."""
    loader = DocumentLoader()
    return loader.load(file_path)